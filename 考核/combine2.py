import re
from docx import Document
import json

# 打开Word文档
doc = Document('book.docx')

# 初始化两个列表，分别存储文本信息和表格信息的JSON数据
extracted_info = []
tables_json = []

# 遍历文档中的段落
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 检查是否是目标标题段落
    if text == '【概述】':
        # 提取【概述】前一行的内容，确保不取空行
        if i - 1 >= 0 and doc.paragraphs[i - 1].text.strip():
            overview_text = re.sub(r'第\d+节', '', doc.paragraphs[i - 1].text.strip())
            extracted_info.append({'id': overview_text})
        
    # 定义一个辅助函数来提取标题下的内容
    def extract_content_after_title(title):
        content = ''
        for j in range(i + 1, len(doc.paragraphs)):
            if '【' in doc.paragraphs[j].text:
                break
            content += doc.paragraphs[j].text.strip() + ' '
        return content.strip()

    # 检查并提取【临床表现】下的内容
    if '【临床表现】' in text or '【中毒表现及处理】' in text:
        extracted_info[-1]['临床表现'] = extract_content_after_title(title='临床表现')
    
    # 检查并提取【诊断要点】下的内容
    if '【诊断要点】' in text:
        extracted_info[-1]['诊断要点'] = extract_content_after_title(title='诊断要点')
    
    # 检查并提取【处理原则】下的内容
    if '【处理原则】' in text:
        extracted_info[-1]['处理原则'] = extract_content_after_title(title='处理原则')

# 遍历文档中的所有表格
for table in doc.tables:
    # 获取第一行所有单元格的文本内容
    first_row_text = [cell.text for cell in table.rows[0].cells]
    
    # 检查第一行的文本中是否包含所有关键字
    if all(keyword in first_row_text for keyword in ["名称", "理化特性", "暴露机会"]):
        # 初始化一个列表来存储转换后的JSON对象
        json_rows = []
        
        # 遍历表格的其余行（跳过已经作为键的第一行）
        for row in table.rows[1:]:
            # 创建一个字典来存储当前行的数据，使用第一行的文本作为键
            row_data = {}
            for cell, key in zip(row.cells, first_row_text):
                row_data[key] = cell.text
            
            # 将当前行的字典数据添加到列表中
            json_rows.append(row_data)
        
        # 将当前表格的所有行数据转换为JSON格式，并添加到总的JSON数据列表中
        tables_json.append(json_rows)

# 将文本信息和表格信息合并到一个JSON对象中
combined_json = {'text_info': extracted_info, 'tables_info': tables_json}

# 将合并后的JSON对象转换为JSON字符串
json_data = json.dumps(combined_json, ensure_ascii=False, indent=4)

# 打印JSON字符串或者将其写入文件
print(json_data)

# 如果需要将JSON数据写入文件
with open('combined2_data.json', 'w', encoding='utf-8') as json_file:
    json_file.write(json_data)
